import streamlit as st
from typing import Dict, Any, Optional, List
import time

class UIComponents:
    """再利用可能なUIコンポーネント"""
    
    @staticmethod
    def render_header():
        """アプリケーションヘッダーを描画"""
        st.title("🎓 SkimMateAgent")
        st.markdown("""
        PDFをアップロードして、AIエージェントによる対話形式で効果的に理解を深めましょう。
        """)
    
    def render_sidebar_settings(self) -> Dict[str, Any]:
        """サイドバーに設定を描画"""
        from services.session_manager import SessionManager

        with st.sidebar:
            # 設定ロック状態を確認
            is_locked = SessionManager.is_settings_locked()

            # ロック時のヘッダー
            if is_locked:
                st.header("🔒 設定（処理中のためロック）")
                st.warning("⚠️ Q&A処理中は設定を変更できません")
            else:
                st.header("⚙️ 設定")

            # 設定コンテナ（ロック時は無効化）
            with st.container():
                # 処理モード設定
                st.subheader("⚡ 処理モード")
                processing_settings = self.render_processing_mode_sidebar(disabled=is_locked)

                st.divider()

                # Q&A設定
                st.subheader("💬 Q&A設定")
                qa_settings = self.render_basic_qa_settings_sidebar(disabled=is_locked)

                st.divider()

                # フォローアップ設定
                st.subheader("🔄 フォローアップ")
                followup_settings = self.render_followup_settings_sidebar(disabled=is_locked)

                st.divider()

                # 重要単語設定
                st.subheader("📝 重要単語設定")
                keyword_settings = self.render_keyword_settings_sidebar(qa_settings.get('qa_turns', 10), disabled=is_locked)

                st.divider()

            # リセットボタンのスタイルを追加
            st.markdown("""
            <style>
            /* サイドバーリセットボタンのスタイル */
            div[data-testid="stSidebar"] div[data-testid="stButton"] > button {
                background: linear-gradient(135deg, rgba(173, 216, 230, 0.4) 0%, rgba(135, 206, 235, 0.5) 50%, rgba(176, 224, 230, 0.4) 100%) !important;
                border: 1px solid rgba(173, 216, 230, 0.6) !important;
                border-radius: 12px !important;
                color: #2c5aa0 !important;
                font-weight: 500 !important;
                backdrop-filter: blur(10px) !important;
                -webkit-backdrop-filter: blur(10px) !important;
                box-shadow:
                    0 2px 8px rgba(173, 216, 230, 0.3),
                    inset 0 1px 0 rgba(255, 255, 255, 0.2) !important;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
            }

            div[data-testid="stSidebar"] div[data-testid="stButton"] > button:hover {
                background: linear-gradient(135deg, rgba(173, 216, 230, 0.6) 0%, rgba(135, 206, 235, 0.7) 50%, rgba(176, 224, 230, 0.6) 100%) !important;
                transform: translateY(-2px) !important;
                box-shadow:
                    0 4px 16px rgba(173, 216, 230, 0.4),
                    inset 0 1px 0 rgba(255, 255, 255, 0.3) !important;
            }
            </style>
            """, unsafe_allow_html=True)

            # リセットボタン（常時表示）
            if st.button("🔄 セッションをリセット", use_container_width=True, help="すべての設定と処理状態を初期化します"):
                from services.session_manager import SessionManager
                # 認証状態を保持
                password_correct = st.session_state.get("password_correct", False)

                # セッション全体をリセット
                for key in list(st.session_state.keys()):
                    del st.session_state[key]

                # 認証状態を復元（ログイン済み状態を維持）
                st.session_state["password_correct"] = password_correct

                SessionManager.initialize_session()
                SessionManager.unlock_settings()  # 設定ロックも解除
                SessionManager.stop_processing()  # 処理状態もリセット
                st.success("✅ セッションをリセットしました")
                st.rerun()

            # ログアウトボタン
            if st.button("🔓 ログアウト", use_container_width=True):
                from auth import logout
                logout()

            # 設定を統合して返す（モデル設定は空の辞書）
            model_settings = {}
            settings = {**processing_settings, **qa_settings, **followup_settings, **keyword_settings, **model_settings}
            return settings


    def render_processing_mode_sidebar(self, disabled: bool = False) -> Dict[str, Any]:
        """サイドバー用処理モード設定を描画"""
        settings = {}

        st.markdown('<div class="sidebar-group">', unsafe_allow_html=True)

        # Quickモード設定をカードデザインで
        settings['quick_mode'] = st.checkbox(
            "Quickモード",
            value=False,
            key="sidebar_quick_mode",
            disabled=disabled,
            help="最終レポートをAI生成せず、要約とQ&Aを単純結合して高速化します"
        )

        # モード説明カード
        if settings['quick_mode']:
            st.markdown("""
            <div style="
                background: linear-gradient(135deg, #FFA50015 0%, #FF8C0008 100%);
                border: 2px solid #FFA50040;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 24px; margin-right: 12px;">💨</span>
                    <span style="
                        font-weight: 700;
                        font-size: 16px;
                        color: #FF8C00;
                    ">Quickモード有効</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">最終レポートは簡易形式で即座に生成されます。処理時間が大幅に短縮されます。</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="
                background: linear-gradient(135deg, #4ECDC415 0%, #45B7D108 100%);
                border: 2px solid #4ECDC440;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 24px; margin-right: 12px;">🎯</span>
                    <span style="
                        font-weight: 700;
                        font-size: 16px;
                        color: #4ECDC4;
                    ">標準モード</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">AIが学習内容を分析して詳細なレポートを生成します。高品質な要約が作成されます。</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)
        return settings

    def render_basic_qa_settings_sidebar(self, disabled: bool = False) -> Dict[str, Any]:
        """サイドバー用基本Q&A設定を描画"""
        settings = {}

        st.markdown('<div class="sidebar-group">', unsafe_allow_html=True)

        # Q&Aターン数設定
        settings['qa_turns'] = st.slider(
            "Q&A数",
            min_value=1,
            max_value=20,
            value=5,
            step=1,
            disabled=disabled,
            key="sidebar_qa_turns",
            help="生成するQ&Aペアの数を設定します"
        )

        # 質問レベル設定
        st.markdown("### 📚 質問レベル")

        level_options = {
            "🔰 簡単": {
                "value": "simple",
                "description": "基本的で理解しやすい質問",
                "detail": "文書の主要なポイントや基本的な仕組みについての質問",
                "color": "#4CAF50"
            },
            "🎓 詳細": {
                "value": "latest",
                "description": "専門的で深い理解を求める質問",
                "detail": "文書の詳細な分析や論理的な関係性についての質問",
                "color": "#2196F3"
            }
        }

        # デフォルトは簡単レベル
        default_level = "🔰 簡単"

        selected_level = st.radio(
            "質問の難易度を選択",
            options=list(level_options.keys()),
            index=0,  # デフォルトで簡単レベル
            disabled=disabled,
            key="sidebar_question_level",
            help="学生エージェントが生成する質問の難易度を選択します"
        )

        settings['question_level'] = level_options[selected_level]["value"]

        # 選択されたレベルの説明を表示
        selected_info = level_options[selected_level]
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, {selected_info['color']}15 0%, {selected_info['color']}08 100%);
            border: 2px solid {selected_info['color']}40;
            border-radius: 12px;
            padding: 15px;
            margin: 10px 0;
        ">
            <div style="color: {selected_info['color']}; font-weight: 600; margin-bottom: 8px;">
                {selected_level}: {selected_info['description']}
            </div>
            <div style="color: #666; font-size: 13px; line-height: 1.4;">
                {selected_info['detail']}
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)
        return settings

    def render_followup_settings_sidebar(self, disabled: bool = False) -> Dict[str, Any]:
        """サイドバー用フォローアップ設定を描画"""
        settings = {}

        settings['enable_followup'] = st.checkbox(
            "追加質問機能を有効化",
            value=False,
            key="sidebar_enable_followup_checkbox",
            help="学習者の理解度に応じて追加の質問を自動生成します",
            disabled=disabled
        )

        if settings['enable_followup']:
            # カスタムスタイルのレベル選択カード
            st.markdown('<div class="sidebar-group">', unsafe_allow_html=True)
            st.markdown("### 📈 追加質問のレベル設定")

            # 3段階のオプション定義
            threshold_options = {
                "🔥 積極的": {
                    "value": 0.2,
                    "description": "基本的な回答でも追加質問を生成",
                    "detail": "簡単な答えでも満足せず、より詳しい専門的な説明を求めます",
                    "color": "#FF6B6B",
                    "icon": "🔥"
                },
                "⚖️ バランス": {
                    "value": 0.5,
                    "description": "適度に専門的な回答で追加質問",
                    "detail": "ある程度詳しい説明があっても、重要な部分はさらに深く掘り下げます",
                    "color": "#4ECDC4",
                    "icon": "⚖️"
                },
                "🎯 厳選": {
                    "value": 0.8,
                    "description": "高度で専門的な回答のみ追加質問",
                    "detail": "かなり詳細で専門的な説明がされた場合のみ、さらに深い質問を生成します",
                    "color": "#45B7D1",
                    "icon": "🎯"
                }
            }

            # カスタムラジオボタンスタイル
            selected_level = st.radio(
                "追加質問の頻度を選択してください",
                options=list(threshold_options.keys()),
                index=1,  # デフォルトは「バランス」
                key="sidebar_followup_level",
                help="回答の専門性に応じた追加質問の発生頻度を調整できます",
                disabled=disabled
            )

            # 選択されたレベルの詳細説明をカードで表示
            selected_config = threshold_options[selected_level]

            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, {selected_config['color']}15 0%, {selected_config['color']}08 100%);
                border: 2px solid {selected_config['color']}40;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 24px; margin-right: 12px;">{selected_config['icon']}</span>
                    <span style="
                        font-weight: 700;
                        font-size: 18px;
                        color: {selected_config['color']};
                    ">{selected_level.split(' ', 1)[1]}</span>
                </div>
                <div style="
                    font-weight: 600;
                    color: var(--neutral-700);
                    margin-bottom: 8px;
                    font-size: 14px;
                ">{selected_config['description']}</div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">{selected_config['detail']}</div>
            </div>
            """, unsafe_allow_html=True)

            # 設定値を保存
            settings['followup_threshold'] = selected_config['value']
            st.markdown('</div>', unsafe_allow_html=True)

            # 最大追加質問数セクション
            st.markdown('<div class="sidebar-group">', unsafe_allow_html=True)
            st.markdown("### 🔢 最大追加質問数")

            settings['max_followups'] = st.slider(
                "1つのトピックにつき最大何問まで追加質問しますか？",
                min_value=1,
                max_value=3,
                value=1,  # デフォルトは1
                step=1,
                key="sidebar_max_followups",
                help="同じトピックについて生成する追加質問の上限数です"
            )

            st.markdown('</div>', unsafe_allow_html=True)

        else:
            settings['followup_threshold'] = 0.5
            settings['max_followups'] = 0
            st.info("💡 追加質問機能を有効にすると、理解が浅い部分について自動的に詳しく質問を生成します")

        return settings

    def render_keyword_settings_sidebar(self, qa_turns: int, disabled: bool = False) -> Dict[str, Any]:
        """サイドバー用重要単語設定を描画"""
        settings = {}

        st.markdown('<div class="sidebar-group">', unsafe_allow_html=True)

        # 機能説明カード
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #6C63FF15 0%, #5A52FF08 100%);
            border: 2px solid #6C63FF40;
            border-radius: 12px;
            padding: 16px;
            margin: 12px 0;
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        ">
            <div style="
                display: flex;
                align-items: center;
                margin-bottom: 8px;
            ">
                <span style="font-size: 20px; margin-right: 8px;">🎯</span>
                <span style="
                    font-weight: 700;
                    font-size: 16px;
                    color: #6C63FF;
                ">重要単語機能</span>
            </div>
            <div style="
                color: var(--neutral-600);
                font-size: 13px;
                line-height: 1.5;
                margin-bottom: 8px;
            ">指定した単語について優先的に質問を生成し、重要なキーワードを確実に学習できます</div>
            <div style="
                color: var(--neutral-500);
                font-size: 12px;
            ">💡 単語数はQ&A数より少なくしてください</div>
        </div>
        """, unsafe_allow_html=True)

        keyword_input = st.text_input(
            "重要単語を入力（カンマ区切り）",
            placeholder="例: 機械学習, ニューラルネットワーク, 深層学習",
            key="sidebar_keyword_input",
            help="これらの単語について優先的に質問が生成されます。複数の単語はカンマで区切ってください。",
            disabled=disabled
        )

        keywords = []
        if keyword_input.strip():
            keywords = [kw.strip() for kw in keyword_input.split(',') if kw.strip()]

        settings['target_keywords'] = keywords

        # ステータスカード
        if keywords and len(keywords) >= qa_turns:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, #FF6B6B15 0%, #FF534308 100%);
                border: 2px solid #FF6B6B40;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 20px; margin-right: 8px;">❌</span>
                    <span style="
                        font-weight: 700;
                        font-size: 16px;
                        color: #FF6B6B;
                    ">設定エラー</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">単語数({len(keywords)})がQ&A数({qa_turns})以上です。単語を{len(keywords) - qa_turns + 1}個減らすか、Q&A数を増やしてください。</div>
            </div>
            """, unsafe_allow_html=True)
        elif keywords and len(keywords) > qa_turns * 0.8:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, #FFA50015 0%, #FF8C0008 100%);
                border: 2px solid #FFA50040;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 20px; margin-right: 8px;">⚠️</span>
                    <span style="
                        font-weight: 700;
                        font-size: 16px;
                        color: #FFA500;
                    ">バランス注意</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">単語数({len(keywords)})がQ&A数({qa_turns})の80%を超えています。他の内容についての質問が少なくなる可能性があります。</div>
            </div>
            """, unsafe_allow_html=True)
        elif keywords:
            keyword_preview = ', '.join(keywords[:3])
            if len(keywords) > 3:
                keyword_preview += f" など{len(keywords)}個"

            recommended_ratio = min(len(keywords) / qa_turns, 0.5)

            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, #4ECDC415 0%, #45B7D108 100%);
                border: 2px solid #4ECDC440;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    justify-content: space-between;
                    margin-bottom: 8px;
                ">
                    <div style="display: flex; align-items: center;">
                        <span style="font-size: 20px; margin-right: 8px;">✅</span>
                        <span style="
                            font-weight: 700;
                            font-size: 16px;
                            color: #4ECDC4;
                        ">設定完了</span>
                    </div>
                    <span style="
                        font-weight: 700;
                        font-size: 14px;
                        color: #4ECDC4;
                        background: rgba(78, 205, 196, 0.2);
                        padding: 4px 8px;
                        border-radius: 6px;
                    ">{len(keywords)}個</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                    margin-bottom: 8px;
                ">登録単語: {keyword_preview}</div>
                <div style="
                    color: var(--neutral-500);
                    font-size: 12px;
                ">📊 全Q&Aの約{recommended_ratio*100:.0f}%が重要単語関連になります</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="
                background: linear-gradient(135deg, #95A5A615 0%, #74858508 100%);
                border: 2px solid #95A5A640;
                border-radius: 12px;
                padding: 16px;
                margin: 12px 0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    margin-bottom: 8px;
                ">
                    <span style="font-size: 20px; margin-right: 8px;">💡</span>
                    <span style="
                        font-weight: 700;
                        font-size: 16px;
                        color: #95A5A6;
                    ">未設定</span>
                </div>
                <div style="
                    color: var(--neutral-600);
                    font-size: 13px;
                    line-height: 1.4;
                ">重要単語未設定。文書全体からバランスよく質問を生成します。</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)
        return settings

    def render_model_settings_sidebar(self) -> Dict[str, Any]:
        """サイドバー用モデル設定を描画"""
        settings = {}

        # GPT-5系のみに制限
        gpt5_models = [
            ('GPT-5', 'gpt-5'),
            ('GPT-5 Mini', 'gpt-5-mini'),
            ('GPT-5 Nano', 'gpt-5-nano')
        ]

        # 推奨モデル設定
        recommended_models = {
            'student': 'gpt-5-mini',
            'teacher': 'gpt-5',
            'summarizer': 'gpt-5-nano'
        }

        # 学生エージェントモデル
        student_model = self._render_model_selector_sidebar(
            "sidebar_student_model",
            gpt5_models,
            recommended_models['student'],
            "🎓 学生モデル"
        )
        settings['student_model'] = student_model

        # 教師エージェントモデル
        teacher_model = self._render_model_selector_sidebar(
            "sidebar_teacher_model",
            gpt5_models,
            recommended_models['teacher'],
            "👨‍🏫 教師モデル"
        )
        settings['teacher_model'] = teacher_model

        # 要約エージェントモデル
        summarizer_model = self._render_model_selector_sidebar(
            "sidebar_summarizer_model",
            gpt5_models,
            recommended_models['summarizer'],
            "📋 要約モデル"
        )
        settings['summarizer_model'] = summarizer_model

        return settings


    def _render_model_selector_sidebar(self, key: str, model_options: list, default_model: str, label: str) -> str:
        """サイドバー用モデル選択セレクトボックスを描画"""
        # セッション状態から現在の値を取得（初回はデフォルト値）
        if key not in st.session_state:
            st.session_state[key] = default_model

        current_value = st.session_state[key]

        # 現在の値に対応するインデックスを取得
        default_index = 0
        for i, (name, model_id) in enumerate(model_options):
            if model_id == current_value:
                default_index = i
                break

        # selectboxで選択されたモデル名を取得
        selected_model_name = st.selectbox(
            label,
            options=[name for name, _ in model_options],
            index=default_index,
            key=f"{key}_selectbox"  # キーの重複を避けるため
        )

        # 選択されたモデルIDを取得してセッション状態を更新
        selected_model_id = next(model_id for name, model_id in model_options if name == selected_model_name)
        st.session_state[key] = selected_model_id

        return selected_model_id

    @staticmethod
    def render_input_options() -> Dict[str, Any]:
        """入力オプション（PDF or テキスト）を描画"""

        # 薄い水色ガラス風タブスタイルを追加
        st.markdown("""
        <style>
        /* Streamlit標準タブのスタイリング */
        div[data-testid="stTabs"] > div[data-baseweb="tab-list"] {
            gap: 8px !important;
            display: flex !important;
        }

        div[data-testid="stTabs"] > div[data-baseweb="tab-list"] > button {
            background: linear-gradient(135deg, rgba(173, 216, 230, 0.4) 0%, rgba(135, 206, 235, 0.5) 50%, rgba(176, 224, 230, 0.4) 100%) !important;
            border: 1px solid rgba(173, 216, 230, 0.6) !important;
            border-radius: 12px !important;
            color: #2c5aa0 !important;
            font-weight: 500 !important;
            padding: 10px 20px !important;
            margin: 0 4px !important;
            backdrop-filter: blur(10px) !important;
            -webkit-backdrop-filter: blur(10px) !important;
            box-shadow:
                0 2px 8px rgba(173, 216, 230, 0.3),
                inset 0 1px 0 rgba(255, 255, 255, 0.2),
                0 1px 3px rgba(0, 0, 0, 0.1) !important;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
            position: relative !important;
        }

        div[data-testid="stTabs"] > div[data-baseweb="tab-list"] > button:hover {
            background: linear-gradient(135deg, rgba(173, 216, 230, 0.6) 0%, rgba(135, 206, 235, 0.7) 50%, rgba(176, 224, 230, 0.6) 100%) !important;
            transform: translateY(-2px) !important;
            box-shadow:
                0 4px 16px rgba(173, 216, 230, 0.4),
                inset 0 1px 0 rgba(255, 255, 255, 0.3),
                0 2px 8px rgba(0, 0, 0, 0.15) !important;
        }

        div[data-testid="stTabs"] > div[data-baseweb="tab-list"] > button[aria-selected="true"] {
            background: linear-gradient(135deg, rgba(173, 216, 230, 0.7) 0%, rgba(135, 206, 235, 0.8) 50%, rgba(176, 224, 230, 0.7) 100%) !important;
            border: 1px solid rgba(135, 206, 235, 0.8) !important;
            color: #1e3a8a !important;
            font-weight: 600 !important;
            transform: translateY(-1px) !important;
            box-shadow:
                0 6px 20px rgba(173, 216, 230, 0.5),
                inset 0 1px 0 rgba(255, 255, 255, 0.4),
                0 3px 12px rgba(0, 0, 0, 0.2) !important;
        }

        /* タブコンテンツのスタイリング */
        div[data-testid="stTabs"] > div[data-baseweb="tab-panel"] {
            padding-top: 20px !important;
            border-top: 1px solid rgba(173, 216, 230, 0.3) !important;
            margin-top: 10px !important;
        }
        </style>
        """, unsafe_allow_html=True)

        # タブで入力方法を選択
        tab1, tab2 = st.tabs(["📁 PDFファイル", "📝 テキスト貼り付け"])

        result = {
            'input_type': None,
            'uploaded_file': None,
            'text_content': None
        }

        with tab1:
            st.markdown("**PDFファイルをアップロード**")
            uploaded_file = st.file_uploader(
                "PDFファイルを選択してください（最大50MB）",
                type=['pdf'],
                help="論文や専門文書のPDFファイルをアップロードしてください",
                key="pdf_uploader"
            )

            if uploaded_file:
                result['input_type'] = 'pdf'
                result['uploaded_file'] = uploaded_file
                st.success(f"✅ PDFファイル: {uploaded_file.name}")

        with tab2:
            st.markdown("**テキストを直接貼り付け**")
            text_content = st.text_area(
                "文書のテキストを貼り付けてください",
                height=200,
                placeholder="ここに文書のテキストを貼り付けてください...\n\n論文、記事、レポートなど、どんなテキストでも分析できます。",
                help="コピー＆ペーストで簡単に文書を入力できます",
                key="text_input"
            )

            if text_content and text_content.strip():
                result['input_type'] = 'text'
                result['text_content'] = text_content.strip()

                # 文字数のみ表示
                char_count = len(text_content)
                st.metric("文字数", f"{char_count:,}")

                if char_count > 500000:  # 50万文字以上で警告
                    st.warning("⚠️ テキストが長すぎる可能性があります。分割して処理することをお勧めします。")
                else:
                    st.success("✅ テキストが入力されました")

        return result
    

    def render_qa_settings(self, disabled: bool = False) -> Dict[str, Any]:
        """Q&A設定を描画"""
        from services.openai_service import OpenAIService
        
        st.subheader("⚙️ Q&A設定")
        
        settings = {}
        
        # Q&Aターン数設定
        settings['qa_turns'] = st.slider(
            "Q&Aターン数",
            min_value=1,
            max_value=20,
            value=10,
            step=1,
            help="生成するQ&Aペアの数を設定してください",
            disabled=disabled
        )
        
        # エージェント別モデル選択
        with st.expander("🤖 エージェント別モデル設定", expanded=False):
            # GPT-5系のみに制限
            gpt5_models = [
                ('GPT-5', 'gpt-5'),
                ('GPT-5 Mini', 'gpt-5-mini'),
                ('GPT-5 Nano', 'gpt-5-nano')
            ]

            # 推奨モデル設定（GPT-5系）
            recommended_models = {
                'student': 'gpt-5-mini',       # 最新軽量モデル
                'teacher': 'gpt-5',            # 最新最高性能モデル
                'summarizer': 'gpt-5-nano'     # 最新超軽量モデル
            }

            # 3つのカラムに分けて表示
            col1, col2, col3 = st.columns(3)

            with col1:
                st.markdown("**🎓 学生エージェント**")
                st.caption("質問生成担当")
                student_model = self._render_model_selector(
                    "student_model",
                    gpt5_models,
                    recommended_models['student'],
                    "質問生成に使用するモデル。軽量モデルでも十分な性能を発揮します。"
                )
                settings['student_model'] = student_model

            with col2:
                st.markdown("**👨‍🏫 教師エージェント**")
                st.caption("回答生成担当")
                teacher_model = self._render_model_selector(
                    "teacher_model",
                    gpt5_models,
                    recommended_models['teacher'],
                    "回答生成に使用するモデル。複雑な内容に対応するため高性能モデルを推奨。"
                )
                settings['teacher_model'] = teacher_model

            with col3:
                st.markdown("**📋 要約エージェント**")
                st.caption("要約・レポート作成担当")
                summarizer_model = self._render_model_selector(
                    "summarizer_model",
                    gpt5_models,
                    recommended_models['summarizer'],
                    "要約とレポート作成に使用するモデル。軽量モデルでも十分な性能を発揮します。"
                    )
                settings['summarizer_model'] = summarizer_model
                    
        # フォローアップ設定
        st.divider()
        st.markdown("**🔄 フォローアップ質問設定**")
        
        settings['enable_followup'] = st.checkbox(
            "フォローアップ質問を有効にする",
            value=True,
            key="enable_followup_checkbox",
            help="回答が専門的すぎる場合に、より理解しやすい説明を求める追加質問を自動生成します"
        )
        
        if settings['enable_followup']:
            # 複雑さ閾値の設定
            settings['followup_threshold'] = st.slider(
                "フォローアップ質問の発動閾値", 
                min_value=0.1,
                max_value=1.0,
                value=0.3,
                step=0.1,
                help="回答の複雑さがこの閾値を超えた場合にフォローアップ質問を生成します"
            )
            
            # 最大フォローアップ数の設定
            settings['max_followups'] = st.slider(
                "最大フォローアップ質問数",
                min_value=0,
                max_value=3,
                value=1,
                step=1,
                help="1つのQ&Aペアに対する最大フォローアップ質問数"
            )
        else:
            # フォローアップ無効時のデフォルト値
            settings['followup_threshold'] = 0.3  # 調整後のデフォルト値
            settings['max_followups'] = 0

        st.divider()
        
        # 単語登録設定
        st.markdown("**📝 重要単語登録**")
        st.caption("指定した単語について必ず質問を生成します（登録単語数 < Q&Aターン数 にしてください）")
        
        # 単語入力
        keyword_input = st.text_input(
            "重要単語を入力（カンマ区切りで複数入力可能）",
            placeholder="例: 機械学習, ニューラルネットワーク, 深層学習",
            help="これらの単語について優先的に質問が生成されます",
            disabled=disabled
        )
        
        # 入力された単語をリストに変換
        keywords = []
        if keyword_input.strip():
            keywords = [kw.strip() for kw in keyword_input.split(',') if kw.strip()]
        
        settings['target_keywords'] = keywords
        
        # 登録単語数とQ&Aターン数の関係をチェック
        qa_turns = settings.get('qa_turns', 10)
        if keywords and len(keywords) >= qa_turns:
            st.warning(f"⚠️ 登録単語数({len(keywords)})がQ&Aターン数({qa_turns})以上です。単語を減らすかターン数を増やしてください。")
        elif keywords:
            st.success(f"✅ {len(keywords)}個の単語を登録: {', '.join(keywords[:3])}{'...' if len(keywords) > 3 else ''}")
        
        return settings
    
    def _render_model_selector(self, key: str, model_options: list, default_model: str, help_text: str) -> str:
        """モデル選択セレクトボックスを描画"""
        # セッション状態に値がある場合はそれを使用、なければデフォルト値
        current_value = st.session_state.get(key, default_model)
        
        # デフォルトモデルのインデックスを取得
        default_index = 0
        for i, (name, model_id) in enumerate(model_options):
            if model_id == current_value:
                default_index = i
                break
        
        selected_model_name = st.selectbox(
            "モデル選択",
            options=[name for name, _ in model_options],
            index=default_index,
            key=key,
            help=help_text
        )
        
        # 選択されたモデルのIDを取得
        selected_model_id = next(model_id for name, model_id in model_options if name == selected_model_name)
        st.caption(f"💡 {selected_model_id}")
        
        return selected_model_id
    
    @staticmethod
    def render_document_info(doc_data: Dict[str, Any]):
        """文書情報を表示"""
        if not doc_data:
            return

        st.subheader("📋 文書情報")

        # テキスト入力とPDF入力で表示を分ける
        source_type = doc_data.get('source_type', 'pdf')

        if source_type == 'text':
            # テキスト入力の場合は文字数のみ表示
            col1, col2 = st.columns(2)

            with col1:
                st.metric("文字数", f"{doc_data.get('char_count', 0):,}")

            with col2:
                split_status = "テキスト入力"
                st.metric("入力形式", split_status)

        else:
            # PDF入力の場合は従来の表示
            col1, col2, col3 = st.columns(3)

            with col1:
                st.metric("ページ数", doc_data.get('page_count', 0))

            with col2:
                st.metric("トークン数", f"{doc_data.get('total_tokens', 0):,}")

            with col3:
                split_status = "分割済み" if doc_data.get('is_split', False) else "未分割"
                st.metric("処理状況", split_status)

            # トークン数による警告
            if doc_data.get('total_tokens', 0) > 200000:
                st.warning("⚠️ トークン数が多いため、処理に時間がかかる可能性があります")
    
    @staticmethod
    def render_progress_indicator(show: bool, text: str):
        """プログレスインジケーターを表示"""
        if show:
            with st.spinner(text):
                time.sleep(0.1)  # UI更新のため短時間待機
    
    @staticmethod
    def render_summary_section(summary: str):
        """要約セクションを描画"""
        if summary:
            st.subheader("📋 文書要約")
            st.markdown(summary)
            st.divider()
    
    @staticmethod
    def render_qa_streaming_area():
        """Q&Aストリーミング表示エリアを作成"""
        st.subheader("💬 Q&Aセッション")
        
        # プレースホルダーを作成
        qa_container = st.container()
        
        return qa_container
    
    @staticmethod
    def render_qa_pair(container, question: str, answer: str, pair_number: int):
        """Q&Aペアを表示"""
        with container:
            with st.expander(f"Q{pair_number}: {question[:50]}...", expanded=True):
                st.markdown(f"**質問：** {question}")
                st.markdown(f"**回答：** {answer}")
    
    @staticmethod
    def render_final_report(report: str):
        """最終レポートを表示"""
        if report:
            st.subheader("📊 最終レポート")
            st.markdown(report)

            # エクスポートボタン
            col1, col2 = st.columns(2)

            with col1:
                # Word形式でダウンロード
                word_content = UIComponents._convert_markdown_to_word_content(report)
                st.download_button(
                    label="📄 Wordでダウンロード",
                    data=word_content,
                    file_name="QA_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="レポートをWord文書形式でダウンロードします"
                )

            with col2:
                # テキスト形式でダウンロード（バックアップ）
                text_content = UIComponents._convert_markdown_to_text(report)
                st.download_button(
                    label="📝 テキストでダウンロード",
                    data=text_content,
                    file_name="QA_Report.txt",
                    mime="text/plain",
                    help="レポートをテキスト形式でダウンロードします"
                )

    @staticmethod
    def _convert_markdown_to_word_content(markdown_text: str) -> bytes:
        """MarkdownテキストをWord文書に変換"""
        try:
            from docx import Document
            from docx.shared import Inches
            import re

            doc = Document()

            # ドキュメントのタイトル
            title = doc.add_heading('AI Q&Aセッション レポート', 0)

            # 現在の日時を追加
            from datetime import datetime
            date_paragraph = doc.add_paragraph(f"生成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
            date_paragraph.runs[0].italic = True

            doc.add_paragraph()  # 空行

            # Markdownを解析してWord文書に変換
            lines = markdown_text.split('\n')
            current_list = None

            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # 空行
                    continue

                # 見出し処理
                if line.startswith('###'):
                    doc.add_heading(line[3:].strip(), level=3)
                elif line.startswith('##'):
                    doc.add_heading(line[2:].strip(), level=2)
                elif line.startswith('#'):
                    doc.add_heading(line[1:].strip(), level=1)

                # リスト処理
                elif line.startswith('- ') or line.startswith('* '):
                    text = line[2:].strip()
                    # 太字処理
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                    doc.add_paragraph(text, style='List Bullet')

                # 通常のテキスト
                else:
                    # 太字処理
                    if '**' in line:
                        p = doc.add_paragraph()
                        parts = re.split(r'\*\*(.*?)\*\*', line)
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                p.add_run(part)
                            else:
                                p.add_run(part).bold = True
                    else:
                        doc.add_paragraph(line)

            # バイト形式で保存
            from io import BytesIO
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            return word_buffer.getvalue()

        except ImportError:
            # python-docxが利用できない場合はテキスト形式で返す
            st.warning("Word形式の出力にはpython-docxライブラリが必要です。テキスト形式でダウンロードしてください。")
            return UIComponents._convert_markdown_to_text(markdown_text).encode('utf-8')
        except Exception as e:
            st.error(f"Word形式への変換エラー: {str(e)}")
            return UIComponents._convert_markdown_to_text(markdown_text).encode('utf-8')

    @staticmethod
    def _convert_markdown_to_text(markdown_text: str) -> str:
        """Markdownテキストをプレーンテキストに変換"""
        import re
        from datetime import datetime

        # ヘッダーを追加
        result = f"AI Q&Aセッション レポート\n"
        result += f"生成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n"
        result += "=" * 50 + "\n\n"

        # Markdownの装飾を除去
        text = re.sub(r'#+\s*', '', markdown_text)  # 見出しマークを除去
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 太字マークを除去
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # イタリックマークを除去
        text = re.sub(r'`(.*?)`', r'\1', text)  # コードマークを除去

        result += text
        return result

    @staticmethod
    def generate_quick_report(summary: str, qa_pairs: list, document_info: dict = None) -> str:
        """Quickモード用の簡易レポートを生成"""
        from datetime import datetime

        # レポートヘッダー
        report = f"""# AI文書要約・Q&Aセッション レポート

**生成日時**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
**処理モード**: Quickモード

---

"""

        # 文書情報
        if document_info:
            report += f"""## 📋 文書情報

- **ページ数**: {document_info.get('page_count', 'N/A')}
- **トークン数**: {document_info.get('total_tokens', 'N/A'):,}
- **Q&A数**: {len(qa_pairs)}

---

"""

        # 要約セクション
        if summary:
            report += f"""## 📄 文書要約

{summary}

---

"""

        # Q&Aセクション
        if qa_pairs:
            report += """## 💬 Q&Aセッション

"""
            for i, qa_pair in enumerate(qa_pairs, 1):
                question = qa_pair.get('question', '質問なし')
                answer = qa_pair.get('answer', '回答なし')

                report += f"""### Q{i}: {question}

**回答**: {answer}

"""

                # フォローアップがある場合
                followup_question = qa_pair.get('followup_question', '')
                followup_answer = qa_pair.get('followup_answer', '')

                if followup_question and followup_answer:
                    report += f"""**追加質問**: {followup_question}

**追加回答**: {followup_answer}

"""

                report += "---\n\n"

        # フッター
        report += f"""## 📊 セッション統計

- **総Q&A数**: {len(qa_pairs)}
- **処理完了**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
- **処理モード**: Quickモード（AI最終レポート生成なし）

*このレポートはQuickモードで生成されました。詳細な分析や洞察が必要な場合は、通常モードをご利用ください。*
"""

        return report
    
    @staticmethod
    def render_statistics(stats: Dict[str, Any]):
        """統計情報を表示"""
        st.subheader("📊 セッション統計")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Q&A数", stats.get('qa_count', 0))
        
        with col2:
            duration = stats.get('duration_seconds', 0)
            duration_str = f"{duration:.1f}秒" if duration else "N/A"
            st.metric("処理時間", duration_str)
        
        with col3:
            st.metric("文書ページ数", stats.get('document_pages', 0))
        
        with col4:
            tokens = stats.get('document_tokens', 0)
            token_str = f"{tokens:,}" if tokens else "N/A"
            st.metric("トークン数", token_str)
    
    @staticmethod
    def render_error_message(error: str):
        """エラーメッセージを表示"""
        st.error(f"❌ エラーが発生しました: {error}")
    
    @staticmethod
    def render_success_message(message: str):
        """成功メッセージを表示"""
        st.success(f"✅ {message}")
    
    @staticmethod
    def render_warning_message(message: str):
        """警告メッセージを表示"""
        st.warning(f"⚠️ {message}")
    
    @staticmethod
    def render_info_message(message: str):
        """情報メッセージを表示"""
        st.info(f"ℹ️ {message}")
    
    @staticmethod
    def render_skeleton_summary():
        """要約セクションのスケルトンローディング"""
        st.subheader("📋 文書要約")
        with st.container():
            # テキストのスケルトン表示
            st.markdown("""
            <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                        background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                        height: 20px; border-radius: 4px; margin: 8px 0;'></div>
            <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                        background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                        height: 20px; border-radius: 4px; margin: 8px 0; width: 85%;'></div>
            <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                        background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                        height: 20px; border-radius: 4px; margin: 8px 0; width: 92%;'></div>
            <style>
            @keyframes shimmer {
              0% { background-position: -200% 0; }
              100% { background-position: 200% 0; }
            }
            </style>
            """, unsafe_allow_html=True)
        st.divider()
    
    @staticmethod
    def render_skeleton_qa():
        """Q&Aセクションのスケルトンローディング"""
        st.subheader("💬 Q&Aセッション")
        for i in range(3):  # 3つのスケルトンQ&Aを表示
            with st.expander(f"❓ Q{i+1}: 質問を生成中...", expanded=False):
                st.markdown("""
                <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                            background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                            height: 16px; border-radius: 4px; margin: 4px 0; width: 70%;'></div>
                <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                            background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                            height: 16px; border-radius: 4px; margin: 8px 0;'></div>
                <div style='background: linear-gradient(90deg, #f0f2f6 25%, #e6e6e6 37%, #f0f2f6 63%); 
                            background-size: 400% 100%; animation: shimmer 1.5s ease-in-out infinite;
                            height: 16px; border-radius: 4px; margin: 4px 0; width: 88%;'></div>
                """, unsafe_allow_html=True)

class StreamingDisplay:
    """ストリーミング表示用のクラス"""
    
    def __init__(self):
        self.containers = {}
        self.current_qa = {"question": "", "answer": ""}
        self.qa_count = 0
    
    def create_streaming_area(self):
        """ストリーミング表示エリアを作成"""
        self.main_container = st.container()
        self.progress_placeholder = st.empty()
        
        return self.main_container
    
    def update_progress(self, text: str):
        """プログレス表示を更新"""
        with self.progress_placeholder:
            st.info(text)
    
    def display_agent_message(self, agent_name: str, content: str, message_type: str = "normal"):
        """エージェントのメッセージを表示"""
        with self.main_container:
            if agent_name == "student":
                self._display_question(content)
            elif agent_name == "teacher":
                self._display_answer(content)
            elif agent_name == "summarizer":
                self._display_summary(content)
    
    def _display_question(self, question: str):
        """質問を表示"""
        self.qa_count += 1
        self.current_qa["question"] = question
        
        with st.expander(f"❓ Q{self.qa_count}: {question[:50]}...", expanded=True):
            st.markdown(f"**質問：** {question}")
            # 回答用のプレースホルダーを作成
            answer_placeholder = st.empty()
            self.current_qa["answer_placeholder"] = answer_placeholder
    
    def _display_answer(self, answer: str):
        """回答を表示"""
        if "answer_placeholder" in self.current_qa:
            with self.current_qa["answer_placeholder"]:
                st.markdown(f"**回答：** {answer}")
        
        # Q&Aペア完了
        self.current_qa = {"question": "", "answer": ""}
    
    def _display_summary(self, summary: str):
        """要約を表示"""
        with self.main_container:
            st.markdown("### 📋 要約")
            st.markdown(summary)
            st.divider()